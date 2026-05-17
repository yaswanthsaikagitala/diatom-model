"""
DiatomAI Guide → Word Document Exporter
Run this script once to generate the Word document.
Requirements: pip install python-docx
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Path to the markdown guide ────────────────────────────────────────────────
GUIDE_PATH = r"C:\Users\Yaswanth\.gemini\antigravity\brain\9cb4efb5-cfcb-4e28-b0ae-c3bde171af5c\diatom_ai_guide.md"
OUTPUT_PATH = r"C:\Users\Yaswanth\Downloads\diatom-classification\DiatomAI_Guide.docx"


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set a table cell's background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Add a thin horizontal line (paragraph border)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def add_code_block(doc, code_text):
    """Add a shaded code block paragraph."""
    p = doc.add_paragraph()
    p.style = "Normal"
    # Grey background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_styled_paragraph(doc, text, bold=False, italic=False,
                          font_size=11, color=None,
                          space_before=0, space_after=4,
                          indent=0):
    """Add a paragraph with inline bold/italic/code rendering."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    _add_inline_runs(p, text, bold=bold, italic=italic,
                     font_size=font_size, color=color)
    return p


def _add_inline_runs(paragraph, text, bold=False, italic=False,
                     font_size=11, color=None):
    """Parse inline **bold**, *italic*, and `code` then add runs."""
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # Text before match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = color
        # The matched segment
        if m.group(2):          # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = italic
            run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = color
        elif m.group(3):        # *italic*
            run = paragraph.add_run(m.group(3))
            run.bold = bold
            run.italic = True
            run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = color
        elif m.group(4):        # `code`
            run = paragraph.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(font_size - 0.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
        pos = m.end()
    # Remaining text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color


def add_table_from_md(doc, header_row, rows):
    """Build a formatted table from parsed markdown rows."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"

    # Header
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header_row):
        hdr_cells[i].text = cell_text.strip()
        set_cell_bg(hdr_cells[i], "2E4057")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    # Data rows
    ROW_COLORS = ["FFFFFF", "F5F7FA"]
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            cells[c_idx].text = cell_text.strip()
            set_cell_bg(cells[c_idx], ROW_COLORS[r_idx % 2])
            for run in cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()   # spacing after table


# ── Main conversion ───────────────────────────────────────────────────────────

def convert(md_path, out_path):
    doc = Document()

    # ── Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ── Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_header = []
    table_rows = []

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()

        # ── Code block fence ─────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────
        if re.match(r'^-{3,}$', stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Tables ───────────────────────────────────────────────────
        if stripped.startswith("|"):
            cells = [c for c in stripped.split("|") if c != ""]
            # separator row (---|---|...)
            if all(re.match(r'^[-:\s]+$', c) for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_header = cells
                table_rows = []
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                add_table_from_md(doc, table_header, table_rows)
                in_table = False
                table_header = []
                table_rows = []

        # ── Headings ─────────────────────────────────────────────────
        h_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if h_match:
            level = len(h_match.group(1))
            heading_text = h_match.group(2)
            # Strip anchor links from heading
            heading_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', heading_text)
            heading_text = re.sub(r'`', '', heading_text)

            size_map = {1: 20, 2: 16, 3: 13, 4: 12}
            color_map = {
                1: RGBColor(0x1A, 0x37, 0x6A),
                2: RGBColor(0x22, 0x55, 0x8A),
                3: RGBColor(0x2E, 0x6D, 0xA0),
                4: RGBColor(0x45, 0x85, 0xAA),
            }
            space_before_map = {1: 16, 2: 14, 3: 10, 4: 8}

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before_map.get(level, 8))
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(size_map.get(level, 11))
            run.font.color.rgb = color_map.get(level, RGBColor(0, 0, 0))
            i += 1
            continue

        # ── Blockquotes (> [!NOTE] style) ────────────────────────────
        if stripped.startswith(">"):
            bq_text = re.sub(r'^>\s*', '', stripped)
            # detect admonition type
            adm_match = re.match(r'^\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]$',
                                 bq_text.strip())
            if adm_match:
                adm_type = adm_match.group(1)
                label_map = {
                    "NOTE": ("ℹ NOTE", "D0E8FF", "1A5276"),
                    "TIP": ("💡 TIP", "D5F5E3", "145A32"),
                    "IMPORTANT": ("❗ IMPORTANT", "FDEBD0", "784212"),
                    "WARNING": ("⚠ WARNING", "FEF9E7", "7D6608"),
                    "CAUTION": ("🚨 CAUTION", "FADBD8", "78281F"),
                }
                label, bg, fg = label_map.get(adm_type,
                                              ("NOTE", "D0E8FF", "1A5276"))
                # Collect continuation lines
                i += 1
                adm_lines = []
                while i < len(lines):
                    nxt = lines[i].rstrip("\n").strip()
                    if nxt.startswith(">"):
                        adm_lines.append(re.sub(r'^>\s*', '', nxt))
                        i += 1
                    else:
                        break
                full_text = label + ": " + " ".join(adm_lines)
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), bg)
                pPr.append(shd)
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(full_text)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(
                    int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16))
                continue
            else:
                # Plain blockquote line
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F5F5F5")
                pPr.append(shd)
                _add_inline_runs(p, bq_text, font_size=10,
                                 color=RGBColor(0x44, 0x44, 0x44))
                i += 1
                continue

        # ── Bullet list ──────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*]\s+(.*)', raw)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            item_text = bullet_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.8 + indent_level * 0.6)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _add_inline_runs(p, item_text, font_size=11)
            i += 1
            continue

        # ── Numbered list ────────────────────────────────────────────
        num_match = re.match(r'^\d+\.\s+(.*)', stripped)
        if num_match:
            item_text = num_match.group(1)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _add_inline_runs(p, item_text, font_size=11)
            i += 1
            continue

        # ── Skip blank lines / table-of-contents links ───────────────
        if stripped == "" or re.match(r'^\d+\.\s+\[.*\]\(#.*\)', stripped):
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────────
        add_styled_paragraph(doc, stripped, font_size=11,
                             space_before=2, space_after=4)
        i += 1

    # flush any remaining table
    if in_table:
        add_table_from_md(doc, table_header, table_rows)

    doc.save(out_path)
    print(f"✅ Word document saved to:\n   {out_path}")


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if not os.path.exists(GUIDE_PATH):
        print(f"❌ Guide not found at:\n   {GUIDE_PATH}")
    else:
        convert(GUIDE_PATH, OUTPUT_PATH)
